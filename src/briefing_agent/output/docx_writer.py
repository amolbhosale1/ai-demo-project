from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path

from docx import Document

from briefing_agent.types import BriefingItem


def write_docx(
    output_dir: str | Path,
    client_name: str,
    summary: str,
    approved_items: list[BriefingItem],
    mix_counts: dict[str, int] | None = None,
) -> Path:
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    now_utc = datetime.now(UTC)
    filename = f"{client_name.lower().replace(' ', '_')}_briefing_{now_utc.date()}.docx"
    file_path = output_path / filename

    doc = Document()
    doc.add_heading(f"Weekly Briefing - {client_name}", level=1)
    doc.add_paragraph(f"Generated: {now_utc.isoformat()} UTC")

    doc.add_heading("Summary", level=2)
    _write_structured_summary(doc, summary)
    if mix_counts:
        doc.add_paragraph(
            f"Source mix - RSS: {mix_counts.get('rss', 0)}, "
            f"API: {mix_counts.get('api', 0)}, "
            f"Other: {mix_counts.get('other', 0)}"
        )

    doc.add_heading("Articles (with links)", level=2)
    for item in approved_items:
        if item.category != "article":
            continue
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item.title + "\n")
        p.add_run(item.url)

    doc.add_heading("Data points", level=2)
    for item in approved_items:
        if item.category != "data_point":
            continue
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item.summary)

    doc.save(file_path)
    return file_path


def _write_structured_summary(doc: Document, summary: str) -> None:
    for raw_line in summary.splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
            continue
        doc.add_paragraph(line)

